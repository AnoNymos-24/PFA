"""
Générateur DOCX — remplit un template Word (.docx) avec des données et insère un QR code.

Fonctionnement :
  1. Remplace tous les marqueurs [champ] par les valeurs fournies.
  2. Détecte la zone rouge (carré QR code) et la remplace par l'image QR générée.

Les marqueurs peuvent se trouver dans :
  - Paragraphes du corps
  - Tableaux (toutes cellules)
  - En-têtes et pieds de page

Gestion des runs fractionnés :
  Word peut découper un marqueur [champ] sur plusieurs runs (selon la mise en forme).
  Ce générateur fusionne d'abord tous les runs d'un paragraphe avant de faire le remplacement,
  puis réapplique le formatage du premier run sur le résultat.
"""

import os
import re
import logging
from pathlib import Path
from typing import Optional, Any

from docx import Document
from docx.shared import Cm
from docx.text.paragraph import Paragraph
from lxml import etree

from app.utils.files import cleanup_temp_file
from .qr_generator import generer_qrcode_fichier_temp

logger = logging.getLogger("smartdoc.docgen.docx")

# ── Namespaces XML ─────────────────────────────────────────────────────────
W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS  = "http://schemas.openxmlformats.org/drawingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
V_NS  = "urn:schemas-microsoft-com:vml"

RED_COLORS = {"ff0000", "c00000", "dc143c", "b22222", "ed0000"}
FIELD_PATTERN = re.compile(r"\[([^\]]+)\]")


# ══════════════════════════════════════════════════════════════════════════
# REMPLACEMENT DES CHAMPS [champ]
# ══════════════════════════════════════════════════════════════════════════

def _remplacer_dans_paragraphe(para: Paragraph, data: dict[str, Any]) -> None:
    """
    Remplace les marqueurs [champ] dans un paragraphe.

    Stratégie :
    - Reconstitue le texte complet depuis tous les runs.
    - Si un marqueur est trouvé, effectue le remplacement.
    - Écrit le résultat dans le premier run et vide les suivants
      (préserve le formatage du premier run).
    """
    if not para.runs:
        return

    texte_total = "".join(run.text for run in para.runs)
    if not FIELD_PATTERN.search(texte_total):
        return

    def _remplacer(match: re.Match) -> str:
        cle = match.group(1).strip()
        valeur = data.get(cle)
        if valeur is None:
            return match.group(0)          # Conserver le marqueur si valeur absente
        return str(valeur)

    nouveau_texte = FIELD_PATTERN.sub(_remplacer, texte_total)

    para.runs[0].text = nouveau_texte
    for run in para.runs[1:]:
        run.text = ""


def _remplacer_champs(doc: Document, data: dict[str, Any]) -> list[str]:
    """
    Remplace tous les marqueurs [champ] dans le document.

    Returns:
        Liste des champs dont la valeur n'était pas fournie (non remplacés).
    """
    # Collecter les champs non remplis
    champs_absents: set[str] = set()
    for match in FIELD_PATTERN.finditer(_extraire_texte_complet(doc)):
        cle = match.group(1).strip()
        if cle not in data:
            champs_absents.add(cle)

    # Corps principal
    for para in doc.paragraphs:
        _remplacer_dans_paragraphe(para, data)

    # Tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _remplacer_dans_paragraphe(para, data)

    # En-têtes et pieds de page
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part:
                for para in part.paragraphs:
                    _remplacer_dans_paragraphe(para, data)

    return sorted(champs_absents)


def _extraire_texte_complet(doc: Document) -> str:
    """Retourne tout le texte du document (pour analyser les champs manquants)."""
    parties = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parties.extend(p.text for p in cell.paragraphs)
    return "\n".join(parties)


# ══════════════════════════════════════════════════════════════════════════
# DÉTECTION ET REMPLACEMENT DE LA ZONE ROUGE (QR CODE)
# ══════════════════════════════════════════════════════════════════════════

def _trouver_run_rouge(doc: Document) -> Optional[tuple]:
    """
    Localise le run XML contenant la zone rouge (DrawingML ou VML).

    Returns:
        (run_element, para_element, taille_cm) ou None si non trouvé.
    """
    body = doc.element.body

    # ── DrawingML (Word 2007+) ────────────────────────────────────────────
    for srgb in body.iter(f"{{{A_NS}}}srgbClr"):
        if srgb.get("val", "").lower() not in RED_COLORS:
            continue

        # Remonter vers le <w:drawing>
        drawing = srgb
        while drawing is not None and drawing.tag != f"{{{W_NS}}}drawing":
            drawing = drawing.getparent()
        if drawing is None:
            continue

        # Lire les dimensions
        taille_cm = 3.0
        ext = drawing.find(f".//{{{WP_NS}}}extent")
        if ext is not None:
            cx = int(ext.get("cx", 2743200))
            cy = int(ext.get("cy", 2743200))
            taille_cm = round(min(cx, cy) / 914400 * 2.54, 2)

        run_elem  = drawing.getparent()
        para_elem = run_elem.getparent() if run_elem is not None else None

        if run_elem is not None and para_elem is not None:
            return run_elem, para_elem, taille_cm

    # ── VML (Word 2003 / anciennes versions) ──────────────────────────────
    for tag in (f"{{{V_NS}}}rect", f"{{{V_NS}}}shape"):
        for shape in body.iter(tag):
            attrs = (shape.get("fillcolor", "") + shape.get("strokecolor", "")).lower()
            if not (any(c in attrs for c in RED_COLORS) or "red" in attrs):
                continue

            # Remonter vers <w:r>
            cur = shape
            run_elem = None
            para_elem = None
            while cur is not None:
                if cur.tag == f"{{{W_NS}}}r":
                    run_elem = cur
                elif cur.tag == f"{{{W_NS}}}p":
                    para_elem = cur
                    break
                cur = cur.getparent()

            if run_elem is not None and para_elem is not None:
                return run_elem, para_elem, 3.0

    return None


def _inserer_qrcode(
    doc: Document,
    qr_data: str,
    taille_cm: float = 3.0,
) -> bool:
    """
    Remplace la zone rouge par un QR code encodant `qr_data`.

    Steps :
      1. Localise le run contenant la zone rouge.
      2. Supprime ce run du paragraphe.
      3. Récupère l'objet Paragraph python-docx correspondant.
      4. Ajoute un run avec l'image QR code.

    Returns:
        True si le QR code a été inséré, False si aucune zone rouge trouvée.
    """
    resultat = _trouver_run_rouge(doc)
    if resultat is None:
        logger.warning("Aucune zone rouge trouvée — QR code non inséré.")
        return False

    run_elem, para_elem, taille_detectee = resultat
    taille_finale = max(taille_detectee, taille_cm)

    # Supprimer le run contenant la zone rouge
    para_elem.remove(run_elem)

    # Retrouver l'objet Paragraph python-docx qui encapsule ce para_elem
    para_obj = _trouver_para_obj(doc, para_elem)
    if para_obj is None:
        logger.error("Paragraphe parent non retrouvé après suppression du run rouge.")
        return False

    # Générer le QR code dans un fichier temporaire
    chemin_qr = generer_qrcode_fichier_temp(qr_data)
    try:
        new_run = para_obj.add_run()
        new_run.add_picture(chemin_qr, width=Cm(taille_finale), height=Cm(taille_finale))
        logger.info(f"QR code inséré ({taille_finale:.1f}×{taille_finale:.1f} cm).")
        return True
    finally:
        cleanup_temp_file(chemin_qr)


def _trouver_para_obj(doc: Document, para_elem) -> Optional[Paragraph]:
    """
    Retrouve l'objet Paragraph python-docx à partir de son élément XML.
    Cherche dans le corps, les tableaux et les en-têtes/pieds de page.
    """
    # Corps principal
    for para in doc.paragraphs:
        if para._p is para_elem:
            return para

    # Tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para._p is para_elem:
                        return para

    # En-têtes et pieds de page
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part:
                for para in part.paragraphs:
                    if para._p is para_elem:
                        return para

    return None


# ══════════════════════════════════════════════════════════════════════════
# POINT D'ENTRÉE PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════

def remplir_template_docx(
    chemin_template: str,
    data: dict[str, Any],
    chemin_sortie: Path,
    qr_data: Optional[str] = None,
    qr_taille_cm: float = 3.0,
) -> list[str]:
    """
    Remplit un template .docx et produit le fichier de sortie.

    Args:
        chemin_template : chemin du fichier .docx modèle
        data            : dictionnaire {champ: valeur}
        chemin_sortie   : chemin du fichier DOCX produit
        qr_data         : données à encoder dans le QR code (None = pas de QR)
        qr_taille_cm    : taille du QR code en cm (si zone détectée plus grande, priorité à celle-ci)

    Returns:
        Liste des champs non remplis (marqueurs conservés tels quels).
    """
    doc = Document(chemin_template)

    # 1. Remplacement des champs [champ]
    champs_non_remplis = _remplacer_champs(doc, data)
    if champs_non_remplis:
        logger.warning(f"Champs non remplis (valeurs manquantes) : {champs_non_remplis}")

    # 2. Insertion du QR code (si demandé)
    if qr_data:
        _inserer_qrcode(doc, qr_data, taille_cm=qr_taille_cm)

    # 3. Sauvegarde
    chemin_sortie.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(chemin_sortie))
    logger.info(f"Document généré : {chemin_sortie} ({chemin_sortie.stat().st_size} octets)")

    return champs_non_remplis
