"""
Service de gestion des modèles de documents.

Fonctionnement :
  • L'utilisateur fournit un fichier Word (.docx) + un id_type_document.
  • Le service analyse le document pour :
      - Extraire les champs à compléter : texte entre crochets [champ]
      - Détecter la zone QR code        : rectangle/carré rouge dans le document
  • Les métadonnées sont persistées dans un fichier JSON (registry.json).
  • Le fichier .docx est stocké localement dans templates_storage/.
"""

import re
import json
import logging
import shutil
import requests
from datetime import datetime
from pathlib import Path
from typing import Optional

from fastapi import HTTPException, UploadFile
from docx import Document
from lxml import etree

from app.config import settings
from app.models.document import DocumentTemplateRecord

logger = logging.getLogger("smartdoc.templates.service")

# ── Namespaces XML OOXML ───────────────────────────────────────────────────
W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS  = "http://schemas.openxmlformats.org/drawingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
V_NS  = "urn:schemas-microsoft-com:vml"

# Couleurs considérées comme « rouge » pour la zone QR
RED_COLORS = {"ff0000", "c00000", "dc143c", "b22222", "ff0000", "ed0000"}


# ══════════════════════════════════════════════════════════════════════════
# ANALYSE DU DOCUMENT
# ══════════════════════════════════════════════════════════════════════════

def _extraire_champs(doc: Document) -> list[str]:
    """
    Scanne tous les paragraphes, tableaux et en-têtes/pieds de page
    pour trouver les marqueurs [champ à compléter].

    Retourne la liste triée et dédoublonnée des noms de champs.
    """
    PATTERN = re.compile(r"\[([^\]]+)\]")
    champs: set[str] = set()

    def _scanner_paragraphes(paragraphes):
        for para in paragraphes:
            for match in PATTERN.finditer(para.text):
                champs.add(match.group(1).strip())

    # Corps principal
    _scanner_paragraphes(doc.paragraphs)

    # Tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _scanner_paragraphes(cell.paragraphs)

    # En-têtes & pieds de page
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part:
                _scanner_paragraphes(part.paragraphs)

    return sorted(champs)


def _detecter_zone_qrcode(doc: Document) -> tuple[bool, float]:
    """
    Détecte la présence d'une zone carrée rouge dans le document Word.
    Cette zone désigne l'emplacement d'insertion du QR code.

    Supporte :
      • DrawingML (Word 2007+) : cherche <a:srgbClr val="FF0000"/>
      • VML (Word 2003)        : cherche <v:rect fillcolor="red"/>

    Retourne (trouvé: bool, taille_cm: float).
    """
    body = doc.element.body

    # ── DrawingML ─────────────────────────────────────────────────────────
    for srgb in body.iter(f"{{{A_NS}}}srgbClr"):
        val = srgb.get("val", "").lower()
        if val in RED_COLORS:
            # Remonter jusqu'au <w:drawing> pour lire les dimensions
            cur = srgb
            while cur is not None:
                if cur.tag == f"{{{W_NS}}}drawing":
                    ext = cur.find(f".//{{{WP_NS}}}extent")
                    if ext is not None:
                        cx = int(ext.get("cx", 2743200))   # 3 cm en EMU
                        cy = int(ext.get("cy", 2743200))
                        taille_cm = round(min(cx, cy) / 914400 * 2.54, 2)
                        logger.info(
                            f"Zone QR code détectée (DrawingML, rouge={val}, "
                            f"{taille_cm:.1f}×{taille_cm:.1f} cm)"
                        )
                        return True, taille_cm
                    return True, 3.0
                cur = cur.getparent()

    # ── VML ───────────────────────────────────────────────────────────────
    for tag in (f"{{{V_NS}}}rect", f"{{{V_NS}}}shape"):
        for shape in body.iter(tag):
            attrs = (
                shape.get("fillcolor", "")
                + shape.get("strokecolor", "")
                + shape.get("style", "")
            ).lower()
            if any(c in attrs for c in RED_COLORS) or "red" in attrs:
                # Essayer de lire la taille depuis le style CSS
                style = shape.get("style", "")
                taille_cm = _lire_taille_vml(style)
                logger.info(f"Zone QR code détectée (VML, {taille_cm:.1f} cm)")
                return True, taille_cm

    return False, 3.0


def _lire_taille_vml(style: str) -> float:
    """Extrait la taille depuis un attribut CSS 'style' VML (en cm)."""
    def _valeur_css(prop: str) -> Optional[float]:
        for item in style.split(";"):
            if prop in item.lower():
                raw = item.split(":")[1].strip() if ":" in item else ""
                for unite, facteur in [("cm", 1.0), ("in", 2.54), ("pt", 0.0353)]:
                    if unite in raw:
                        try:
                            return float(raw.replace(unite, "").strip()) * facteur
                        except ValueError:
                            pass
        return None

    w = _valeur_css("width")
    h = _valeur_css("height")
    if w and h:
        return round(min(w, h), 2)
    return 3.0


# ══════════════════════════════════════════════════════════════════════════
# REGISTRE (persistance JSON)
# ══════════════════════════════════════════════════════════════════════════

def _charger_registre() -> dict:
    """Charge le registre JSON depuis le disque."""
    f = settings.TEMPLATES_REGISTRY_FILE
    if f.exists():
        with f.open(encoding="utf-8") as fp:
            return json.load(fp)
    return {"next_id": 1, "templates": {}}


def _sauvegarder_registre(registre: dict) -> None:
    """Sauvegarde le registre JSON sur le disque."""
    settings.TEMPLATES_STORAGE_DIR.mkdir(parents=True, exist_ok=True)
    with settings.TEMPLATES_REGISTRY_FILE.open("w", encoding="utf-8") as fp:
        json.dump(registre, fp, ensure_ascii=False, indent=2)


# ══════════════════════════════════════════════════════════════════════════
# CHARGEMENT DU FICHIER SOURCE
# ══════════════════════════════════════════════════════════════════════════

def _telecharger_depuis_url(url: str, destination: Path) -> None:
    """Télécharge un fichier depuis une URL HTTP/HTTPS."""
    try:
        resp = requests.get(url, timeout=30, stream=True)
        resp.raise_for_status()
        with destination.open("wb") as f:
            for chunk in resp.iter_content(chunk_size=8192):
                f.write(chunk)
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=f"Impossible de télécharger le fichier depuis l'URL : {exc}",
        )


def _obtenir_fichier_docx(
    url_fichier_modele: str,
    id_modele: int,
) -> Path:
    """
    Récupère et stocke localement le fichier .docx source.
    Supports : URL HTTP/HTTPS, chemin local absolu/relatif.

    Retourne le chemin local du fichier stocké.
    """
    settings.TEMPLATES_STORAGE_DIR.mkdir(parents=True, exist_ok=True)
    destination = settings.TEMPLATES_STORAGE_DIR / f"template_{id_modele}.docx"

    if url_fichier_modele.startswith(("http://", "https://")):
        _telecharger_depuis_url(url_fichier_modele, destination)
    else:
        src = Path(url_fichier_modele)
        if not src.exists():
            raise HTTPException(
                status_code=400,
                detail=f"Fichier source introuvable : {url_fichier_modele}",
            )
        shutil.copy2(src, destination)

    return destination


# ══════════════════════════════════════════════════════════════════════════
# API DU SERVICE
# ══════════════════════════════════════════════════════════════════════════

async def creer_modele_depuis_upload(
    fichier: UploadFile,
    id_type_document: int,
) -> DocumentTemplateRecord:
    """
    Crée un modèle de document depuis un fichier .docx uploadé.

    Pipeline :
      1. Validation (extension .docx)
      2. Attribution d'un id_modele auto-incrémenté
      3. Sauvegarde du fichier dans templates_storage/
      4. Analyse : extraction des [champs] + détection zone rouge QR
      5. Persistance dans registry.json

    Returns:
        DocumentTemplateRecord avec toutes les métadonnées.
    """
    # ── Validation ────────────────────────────────────────────────────────
    filename = fichier.filename or ""
    if not filename.lower().endswith((".docx", ".doc")):
        raise HTTPException(
            status_code=400,
            detail="Le fichier modèle doit être un document Word (.docx).",
        )

    # ── Attribution ID ────────────────────────────────────────────────────
    registre = _charger_registre()
    id_modele: int = registre["next_id"]
    registre["next_id"] += 1

    # ── Sauvegarde du fichier ─────────────────────────────────────────────
    settings.TEMPLATES_STORAGE_DIR.mkdir(parents=True, exist_ok=True)
    chemin_local = settings.TEMPLATES_STORAGE_DIR / f"template_{id_modele}.docx"

    contenu = await fichier.read()
    with chemin_local.open("wb") as f:
        f.write(contenu)

    logger.info(f"Fichier template sauvegardé : {chemin_local} ({len(contenu)} octets)")

    # ── Analyse du document ───────────────────────────────────────────────
    try:
        doc = Document(str(chemin_local))
    except Exception as exc:
        chemin_local.unlink(missing_ok=True)
        raise HTTPException(
            status_code=422,
            detail=f"Impossible d'ouvrir le fichier Word : {exc}",
        )

    champs     = _extraire_champs(doc)
    a_qr, taille_qr = _detecter_zone_qrcode(doc)

    logger.info(
        f"Analyse template {id_modele} : {len(champs)} champ(s) détecté(s), "
        f"zone QR = {a_qr}"
    )

    # ── URL d'accès au fichier ────────────────────────────────────────────
    url_fichier = f"/api/templates/files/template_{id_modele}.docx"

    # ── Persistance ───────────────────────────────────────────────────────
    record = DocumentTemplateRecord(
        id_modele_document=id_modele,
        url_fichier_modele=url_fichier,
        id_type_document=id_type_document,
        fichier_local=str(chemin_local),
        champs_detectes=champs,
        a_zone_qrcode=a_qr,
        qrcode_taille_cm=taille_qr,
        date_creation=datetime.now().isoformat(),
    )

    registre["templates"][str(id_modele)] = record.model_dump()
    _sauvegarder_registre(registre)

    return record


async def creer_modele_depuis_url(
    url_fichier_modele: str,
    id_type_document: int,
) -> DocumentTemplateRecord:
    """
    Crée un modèle depuis une URL (téléchargement automatique) ou un chemin local.
    Même pipeline que `creer_modele_depuis_upload`.
    """
    registre = _charger_registre()
    id_modele: int = registre["next_id"]
    registre["next_id"] += 1

    chemin_local = _obtenir_fichier_docx(url_fichier_modele, id_modele)

    try:
        doc = Document(str(chemin_local))
    except Exception as exc:
        chemin_local.unlink(missing_ok=True)
        raise HTTPException(status_code=422, detail=f"Fichier Word invalide : {exc}")

    champs     = _extraire_champs(doc)
    a_qr, taille_qr = _detecter_zone_qrcode(doc)

    url_fichier = f"/api/templates/files/template_{id_modele}.docx"

    record = DocumentTemplateRecord(
        id_modele_document=id_modele,
        url_fichier_modele=url_fichier,
        id_type_document=id_type_document,
        fichier_local=str(chemin_local),
        champs_detectes=champs,
        a_zone_qrcode=a_qr,
        qrcode_taille_cm=taille_qr,
        date_creation=datetime.now().isoformat(),
    )

    registre["templates"][str(id_modele)] = record.model_dump()
    _sauvegarder_registre(registre)

    return record


def lister_modeles(id_type: Optional[int] = None) -> list[DocumentTemplateRecord]:
    """Retourne tous les modèles (filtrés par type si fourni)."""
    registre = _charger_registre()
    modeles = [
        DocumentTemplateRecord(**v)
        for v in registre["templates"].values()
    ]
    if id_type is not None:
        modeles = [m for m in modeles if m.id_type_document == id_type]
    return sorted(modeles, key=lambda m: m.id_modele_document)


def obtenir_modele(id_modele: int) -> Optional[DocumentTemplateRecord]:
    """Retourne un modèle par son ID, ou None s'il est introuvable."""
    registre = _charger_registre()
    data = registre["templates"].get(str(id_modele))
    return DocumentTemplateRecord(**data) if data else None


def supprimer_modele(id_modele: int) -> bool:
    """Supprime un modèle et son fichier .docx associé. Retourne True si supprimé."""
    registre = _charger_registre()
    key = str(id_modele)
    if key not in registre["templates"]:
        return False

    # Supprimer le fichier .docx
    chemin = Path(registre["templates"][key]["fichier_local"])
    chemin.unlink(missing_ok=True)

    del registre["templates"][key]
    _sauvegarder_registre(registre)
    logger.info(f"Modèle {id_modele} supprimé.")
    return True
