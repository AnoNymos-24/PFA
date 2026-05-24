"""
Service de génération de documents.

Fonctionnement :
  1. Charger le DocumentTemplateRecord depuis registry.json
     (via id_modele_document)
  2. Ouvrir le fichier .docx template avec python-docx
  3. Remplir tous les marqueurs [champ] avec les données fournies
  4. Si a_zone_qrcode=True et qr_data fourni → générer QR PNG et l'insérer
     à la place de la zone rouge dans le document
  5. Sauvegarder le document final en .docx
  6. Convertir en PDF via LibreOffice si output_format="pdf" (optionnel)
  7. Retourner les métadonnées du fichier généré

Vérification d'authenticité (endpoint /documents/verifier) :
  • Signature HMAC-SHA256 embarquée dans le QR code par le backend Java
  • La vérification ici sert aux documents anciennement générés avec signature
"""

import hashlib
import hmac
import logging
import re
import subprocess
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import Cm

from app.modules.document_templates.service import obtenir_modele
from config import settings

logger = logging.getLogger("smartintern.document_generation.service")

# ── Namespaces XML OOXML ───────────────────────────────────────────────────
W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS  = "http://schemas.openxmlformats.org/drawingml/2006/main"
V_NS  = "urn:schemas-microsoft-com:vml"

RED_COLORS      = {"ff0000", "c00000", "dc143c", "b22222", "ed0000", "cc0000"}
PATTERN_CHAMP   = re.compile(r"\[([^\]]+)\]")


# ══════════════════════════════════════════════════════════════════════════
# REMPLISSAGE DES CHAMPS [champ]
# ══════════════════════════════════════════════════════════════════════════

def _remplir_paragraphe(para, data: dict[str, str]) -> list[str]:
    """
    Remplace les marqueurs [champ] dans un paragraphe Word.

    Stratégie en deux temps :
      1. Remplacement run par run (préserve le formatage gras/italique/couleur)
      2. Si des marqueurs persistent (répartis sur plusieurs runs),
         reconstruction du texte complet dans le premier run

    Retourne la liste des marqueurs toujours présents après remplacement
    (i.e. les champs dont la clé n'est pas dans `data`).
    """
    if not para.text:
        return []

    # Passe 1 : run par run
    for run in para.runs:
        t = run.text
        for key, value in data.items():
            placeholder = f"[{key}]"
            if placeholder in t:
                t = t.replace(placeholder, value)
        run.text = t

    # Passe 2 : si des marqueurs persistent (champ fragmenté sur plusieurs runs)
    remaining = para.text
    if PATTERN_CHAMP.search(remaining):
        full = remaining
        for key, value in data.items():
            full = full.replace(f"[{key}]", value)
        if full != remaining and para.runs:
            para.runs[0].text = full
            for run in para.runs[1:]:
                run.text = ""

    # Collecte des champs toujours non remplis
    return [m.group(1).strip() for m in PATTERN_CHAMP.finditer(para.text)]


def _remplir_champs(doc: Document, data: dict[str, Any]) -> list[str]:
    """
    Parcourt corps + tableaux + en-têtes/pieds et remplace tous les [champs].
    Retourne la liste triée et dédoublonnée des champs non remplis.
    """
    # Convertir toutes les valeurs en str (les None deviennent "")
    data_str = {k: str(v) if v is not None else "" for k, v in data.items()}
    non_remplis: set[str] = set()

    def _parcourir(paragraphes):
        for para in paragraphes:
            non_remplis.update(_remplir_paragraphe(para, data_str))

    _parcourir(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _parcourir(cell.paragraphs)

    for section in doc.sections:
        for part in (section.header, section.footer):
            if part:
                _parcourir(part.paragraphs)

    return sorted(non_remplis)


# ══════════════════════════════════════════════════════════════════════════
# GÉNÉRATION ET INSERTION DU QR CODE
# ══════════════════════════════════════════════════════════════════════════

def _generer_image_qr(qr_data: str, chemin_sortie: str) -> bool:
    """
    Génère une image PNG du QR code à partir des données fournies.
    Retourne True si la génération a réussi.
    """
    try:
        import qrcode  # type: ignore
        qr = qrcode.QRCode(
            version=None,
            error_correction=qrcode.constants.ERROR_CORRECT_H,
            box_size=10,
            border=2,
        )
        qr.add_data(qr_data)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        img.save(chemin_sortie)
        logger.info(f"QR code généré : {chemin_sortie}")
        return True
    except Exception as exc:
        logger.warning(f"Impossible de générer le QR code : {exc}")
        return False


def _inserer_qr_drawingml(body, chemin_qr: str, taille_cm: float, doc: Document) -> bool:
    """
    Tente de remplacer la zone rouge DrawingML par l'image QR.
    Retourne True si réussi.
    """
    for srgb in body.iter(f"{{{A_NS}}}srgbClr"):
        val = srgb.get("val", "").lower()
        if val not in RED_COLORS:
            continue

        # Remonter jusqu'au <w:drawing>
        cur = srgb
        drawing_elem = None
        while cur is not None:
            if cur.tag == f"{{{W_NS}}}drawing":
                drawing_elem = cur
                break
            cur = cur.getparent()
        if drawing_elem is None:
            continue

        # Remonter jusqu'au <w:p>
        cur = drawing_elem.getparent()
        para_elem = None
        while cur is not None:
            if cur.tag == f"{{{W_NS}}}p":
                para_elem = cur
                break
            cur = cur.getparent()
        if para_elem is None:
            continue

        try:
            para_idx = list(body).index(para_elem)
        except ValueError:
            continue

        # Ajouter un paragraphe temporaire avec l'image QR à la fin du doc
        temp_para = doc.add_paragraph()
        run = temp_para.add_run()
        try:
            run.add_picture(chemin_qr, width=Cm(taille_cm), height=Cm(taille_cm))
        except Exception as exc:
            logger.warning(f"Impossible d'insérer l'image QR (DrawingML) : {exc}")
            body.remove(temp_para._element)
            return False

        # Déplacer le paragraphe image avant le placeholder
        img_elem = temp_para._element
        body.remove(img_elem)
        body.insert(para_idx, img_elem)

        # Supprimer le drawing rouge du run parent
        run_parent = drawing_elem.getparent()
        if run_parent is not None:
            run_parent.remove(drawing_elem)

        logger.info(f"QR inséré (DrawingML) à l'index de paragraphe {para_idx}")
        return True

    return False


def _inserer_qr_vml(body, chemin_qr: str, taille_cm: float, doc: Document) -> bool:
    """
    Tente de remplacer la zone rouge VML par l'image QR.
    Retourne True si réussi.
    """
    for tag in (f"{{{V_NS}}}rect", f"{{{V_NS}}}shape"):
        for shape in body.iter(tag):
            attrs = (
                shape.get("fillcolor", "")
                + shape.get("strokecolor", "")
                + shape.get("style", "")
            ).lower()
            if not (any(c in attrs for c in RED_COLORS) or "red" in attrs):
                continue

            # Remonter jusqu'au <w:p>
            cur = shape.getparent()
            para_elem = None
            while cur is not None:
                if cur.tag == f"{{{W_NS}}}p":
                    para_elem = cur
                    break
                cur = cur.getparent()
            if para_elem is None:
                continue

            try:
                para_idx = list(body).index(para_elem)
            except ValueError:
                continue

            temp_para = doc.add_paragraph()
            run = temp_para.add_run()
            try:
                run.add_picture(chemin_qr, width=Cm(taille_cm), height=Cm(taille_cm))
            except Exception as exc:
                logger.warning(f"Impossible d'insérer l'image QR (VML) : {exc}")
                body.remove(temp_para._element)
                return False

            img_elem = temp_para._element
            body.remove(img_elem)
            body.insert(para_idx, img_elem)

            shape_parent = shape.getparent()
            if shape_parent is not None:
                shape_parent.remove(shape)

            logger.info(f"QR inséré (VML) à l'index de paragraphe {para_idx}")
            return True

    return False


def _inserer_qr_dans_doc(doc: Document, chemin_qr: str, taille_cm: float) -> bool:
    """
    Remplace la zone rouge placeholder (DrawingML ou VML) par l'image QR code.
    Retourne True si le remplacement a été effectué.
    """
    body = doc.element.body

    if _inserer_qr_drawingml(body, chemin_qr, taille_cm, doc):
        return True
    if _inserer_qr_vml(body, chemin_qr, taille_cm, doc):
        return True

    logger.warning("Aucune zone rouge (QR placeholder) trouvée dans le document")
    return False


# ══════════════════════════════════════════════════════════════════════════
# CONVERSION PDF
# ══════════════════════════════════════════════════════════════════════════

def _convertir_en_pdf(chemin_docx: Path) -> Optional[Path]:
    """
    Convertit un .docx en PDF via LibreOffice (headless).
    Retourne le chemin du PDF généré, ou None si LibreOffice est absent
    ou si la conversion échoue.
    """
    import platform

    # Localiser l'exécutable LibreOffice
    if platform.system() == "Windows":
        candidats_path = [
            Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
            Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
        ]
        soffice = next((str(p) for p in candidats_path if p.exists()), None)
    else:
        soffice = None
        for candidat in ("libreoffice", "soffice"):
            try:
                result = subprocess.run(
                    ["which", candidat], capture_output=True, text=True, timeout=5
                )
                if result.returncode == 0 and result.stdout.strip():
                    soffice = candidat
                    break
            except Exception:
                pass

    if not soffice:
        logger.warning("LibreOffice introuvable — conversion PDF impossible")
        return None

    try:
        outdir = chemin_docx.parent
        result = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(outdir),
                str(chemin_docx),
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if result.returncode == 0:
            chemin_pdf = chemin_docx.with_suffix(".pdf")
            if chemin_pdf.exists():
                logger.info(f"PDF converti : {chemin_pdf}")
                return chemin_pdf
        logger.error(f"LibreOffice a retourné une erreur : {result.stderr}")
    except subprocess.TimeoutExpired:
        logger.error("Timeout lors de la conversion PDF LibreOffice (120 s)")
    except Exception as exc:
        logger.error(f"Erreur inattendue lors de la conversion PDF : {exc}")

    return None


# ══════════════════════════════════════════════════════════════════════════
# SERVICE PRINCIPAL — GÉNÉRATION
# ══════════════════════════════════════════════════════════════════════════

def generer_document(
    id_modele_document: int,
    data: dict[str, Any],
    qr_data: Optional[str] = None,
    output_format: str = "docx",
    filename: Optional[str] = None,
    doc_id: Optional[str] = None,
) -> tuple[dict, Path]:
    """
    Génère un document personnalisé à partir d'un modèle Word enregistré.

    Pipeline :
      1. Charger le DocumentTemplateRecord depuis registry.json
      2. Ouvrir le .docx template avec python-docx
      3. Remplir les marqueurs [champ] avec les valeurs de `data`
      4. Générer et insérer le QR code dans la zone rouge (si applicable)
      5. Sauvegarder en .docx
      6. Convertir en PDF via LibreOffice si output_format="pdf"
      7. Nettoyer les fichiers temporaires
      8. Retourner (metadata_dict, chemin_fichier_final)

    Args:
        id_modele_document : ID du modèle dans le registry
        data               : dict {nom_champ: valeur} pour remplir les [...]
        qr_data            : texte/URL à encoder dans le QR (optionnel)
        output_format      : "docx" (défaut) ou "pdf" (nécessite LibreOffice)
        filename           : nom personnalisé du fichier généré (sans extension)

    Returns:
        (metadata, chemin_fichier_final)

    Raises:
        ValueError      si le modèle n'existe pas dans le registry
        FileNotFoundError si le fichier .docx du modèle est absent du disque
    """
    # ── 1. Charger le modèle ────────────────────────────────────────────────
    template = obtenir_modele(id_modele_document)
    if template is None:
        raise ValueError(
            f"Modèle #{id_modele_document} introuvable dans le registry."
        )

    chemin_template = Path(template.fichier_local)
    if not chemin_template.exists():
        raise FileNotFoundError(
            f"Fichier modèle introuvable : {chemin_template}. "
            "Il a peut-être été supprimé manuellement."
        )

    # ── 2. Préparer les chemins de sortie ───────────────────────────────────
    # Utiliser le doc_id fourni par le backend (cohérence UUID lien d'auth ↔ fichier)
    doc_id = doc_id or str(uuid.uuid4())
    settings.DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)

    # Nettoyer le nom de fichier fourni (caractères non-alphanumériques → _)
    nom_base = filename or f"document_{doc_id[:8]}"
    nom_base = re.sub(r"[^\w\-_. ]", "_", nom_base).strip()

    chemin_docx = settings.DOCUMENTS_DIR / f"{doc_id}.docx"

    # ── 3. Ouvrir le template et remplir les champs ────────────────────────
    doc = Document(str(chemin_template))
    champs_non_remplis = _remplir_champs(doc, data)

    if champs_non_remplis:
        logger.warning(
            f"Document {doc_id} — champs non remplis : {champs_non_remplis}"
        )

    # ── 4. QR code : générer et insérer ────────────────────────────────────
    qr_insere = False
    chemin_qr: Optional[Path] = None

    if template.a_zone_qrcode and qr_data:
        chemin_qr = settings.DOCUMENTS_DIR / f"{doc_id}_qr.png"
        if _generer_image_qr(qr_data, str(chemin_qr)):
            qr_insere = _inserer_qr_dans_doc(
                doc, str(chemin_qr), template.qrcode_taille_cm
            )
        if not qr_insere:
            logger.warning(f"QR code non inséré pour le document {doc_id}")

    # ── 5. Sauvegarder le .docx rempli ─────────────────────────────────────
    doc.save(str(chemin_docx))
    logger.info(f"Document .docx sauvegardé : {chemin_docx}")

    # ── 6. Conversion PDF (optionnelle) ─────────────────────────────────────
    chemin_final = chemin_docx
    format_final = "docx"

    if output_format == "pdf":
        chemin_pdf = _convertir_en_pdf(chemin_docx)
        if chemin_pdf:
            chemin_final = chemin_pdf
            format_final = "pdf"
            chemin_docx.unlink(missing_ok=True)   # Supprimer le .docx intermédiaire
        else:
            logger.warning(
                f"Conversion PDF échouée pour {doc_id} — "
                "le document .docx est retourné à la place"
            )

    # ── 7. Nettoyage QR temporaire ──────────────────────────────────────────
    if chemin_qr and chemin_qr.exists():
        chemin_qr.unlink(missing_ok=True)

    taille = chemin_final.stat().st_size if chemin_final.exists() else 0

    # ── 8. Métadonnées du document généré ───────────────────────────────────
    metadata = {
        "document_id":        doc_id,
        "filename":           f"{nom_base}.{format_final}",
        "format":             format_final,
        "size_bytes":         taille,
        "download_url":       f"/documents/telecharger/{chemin_final.name}",
        "id_modele_document": id_modele_document,
        "champs_non_remplis": champs_non_remplis,
        "qr_insere":          qr_insere,
    }

    logger.info(
        f"Document généré : {doc_id} | format={format_final} | "
        f"{taille / 1024:.1f} KB | "
        f"{len(champs_non_remplis)} champs manquant(s) | "
        f"QR={'oui' if qr_insere else 'non'}"
    )

    return metadata, chemin_final


# ══════════════════════════════════════════════════════════════════════════
# VÉRIFICATION D'AUTHENTICITÉ (QR code — héritage)
# ══════════════════════════════════════════════════════════════════════════

def _generer_signature(doc_uuid: str, date_generation: str, date_expiration: str) -> str:
    """Signature HMAC-SHA256 pour authentifier un document."""
    message = f"{doc_uuid}:{date_generation}:{date_expiration}"
    return hmac.new(
        settings.SIGNATURE_SECRET.encode(),
        message.encode(),
        hashlib.sha256,
    ).hexdigest()


def _verifier_signature(
    doc_uuid: str, date_generation: str, date_expiration: str, signature: str
) -> bool:
    expected = _generer_signature(doc_uuid, date_generation, date_expiration)
    return hmac.compare_digest(expected, signature)


def _est_expire(date_expiration_str: str) -> bool:
    try:
        return datetime.now() > datetime.fromisoformat(date_expiration_str)
    except Exception:
        return False


def verifier_document(
    doc_uuid: str,
    date_generation: str,
    date_expiration: str,
    signature: str,
    type_document: str = "",
    nom_etablissement: str = "",
) -> dict:
    """
    Vérifie l'authenticité d'un document via sa signature HMAC embarquée dans le QR code.

    Retourne un dict compatible avec ResultatVerification (CamelModel côté Java).
    """
    signature_valide = _verifier_signature(
        doc_uuid, date_generation, date_expiration, signature
    )
    expire          = _est_expire(date_expiration)
    fichier_existe  = (settings.DOCUMENTS_DIR / f"{doc_uuid}.pdf").exists() or \
                      (settings.DOCUMENTS_DIR / f"{doc_uuid}.docx").exists()

    if not signature_valide:
        statut  = "FALSIFIE"
        message = "Document non authentique — signature invalide"
    elif expire:
        statut  = "EXPIRE"
        message = f"Document expiré le {date_expiration[:10]}"
    elif not fichier_existe:
        statut  = "INTROUVABLE"
        message = "Document introuvable dans le système"
    else:
        statut  = "VALIDE"
        message = "Document authentique et valide"

    return {
        "doc_uuid":          doc_uuid,
        "statut":            statut,
        "valide":            statut == "VALIDE",
        "message":           message,
        "type_document":     type_document,
        "nom_etablissement": nom_etablissement,
        "date_generation":   date_generation,
        "date_expiration":   date_expiration,
        "signature_valide":  signature_valide,
        "expire":            expire,
        "fichier_existe":    fichier_existe,
        "verifie_le":        datetime.now().isoformat(),
    }
