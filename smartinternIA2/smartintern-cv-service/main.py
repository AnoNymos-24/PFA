"""
SmartIntern AI — CV Extraction & Scoring Microservice v2.1
FastAPI + pdfplumber + PyMuPDF + pytesseract + Claude AI

v2.1 :
  - Système de fallback IA : provider primaire + provider de secours
  - Basculement automatique si quota/tokens/rate-limit insuffisants
  - Anthropic direct (primaire) + OpenRouter (fallback) par défaut
  - Modèles configurables indépendamment par provider
"""

import os
import re
import json
import tempfile
import subprocess
import logging
from typing import Optional, Any, Annotated
from pathlib import Path
from datetime import datetime

from dotenv import load_dotenv
load_dotenv()

import pdfplumber
import fitz
import pytesseract
from PIL import Image
import openai
from openai import OpenAI
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field, ConfigDict, field_validator, model_validator, BeforeValidator
from pydantic.alias_generators import to_camel

# ── Logging ────────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO,
                    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s")
logger = logging.getLogger("cv-service")

# ══════════════════════════════════════════════════════════════════════════
# CONFIGURATION MULTI-PROVIDER AVEC FALLBACK
#
# Provider 1 (PRIMAIRE) : NVIDIA NIM
#   API key : NVIDIA_API_KEY          (préfixe nvapi-)
#   Base URL: https://integrate.api.nvidia.com/v1
#   Modèle  : NVIDIA_MODEL (défaut : google/gemma-4-31b-it)
#   Max tokens supportés : 16 384
#
# Provider 2 (FALLBACK) : OpenRouter
#   API key : OPENROUTER_API_KEY      (préfixe sk-or-v1-)
#   Base URL: https://openrouter.ai/api/v1
#   Modèle  : OPENROUTER_MODEL (défaut : anthropic/claude-sonnet-4-5)
#
# Basculement automatique si NVIDIA retourne :
#   429 rate-limit · 402 crédits · 503/529 surcharge · context_length_exceeded
# Pour inverser primaire/fallback, échangez les variables dans .env.
# ══════════════════════════════════════════════════════════════════════════

NVIDIA_API_KEY     = os.environ.get("NVIDIA_API_KEY",     "")
NVIDIA_MODEL       = os.environ.get("NVIDIA_MODEL",       "google/gemma-4-31b-it")
NVIDIA_MAX_TOKENS  = int(os.environ.get("NVIDIA_MAX_TOKENS",  "4096"))

OPENROUTER_API_KEY    = os.environ.get("OPENROUTER_API_KEY",    "")
OPENROUTER_MODEL      = os.environ.get("OPENROUTER_MODEL",      "anthropic/claude-sonnet-4-5")
# Plafond pour le compte gratuit OpenRouter (évite l'erreur 402 crédits insuffisants)
OPENROUTER_MAX_TOKENS = int(os.environ.get("OPENROUTER_MAX_TOKENS", "1500"))

if not NVIDIA_API_KEY and not OPENROUTER_API_KEY:
    logger.error("Aucune clé API IA configurée — renseignez NVIDIA_API_KEY et/ou OPENROUTER_API_KEY dans .env")
elif not NVIDIA_API_KEY:
    logger.warning("NVIDIA_API_KEY absente — OpenRouter utilisé seul (pas de fallback)")
elif not OPENROUTER_API_KEY:
    logger.warning("OPENROUTER_API_KEY absente — NVIDIA NIM utilisé seul (pas de fallback)")
else:
    logger.info("Deux providers IA configurés : NVIDIA NIM/Gemma (primaire) + OpenRouter (fallback)")


# ── Codes / mots-clés déclenchant le basculement vers le provider de secours ──

_FALLBACK_STATUS_CODES = {402, 429, 529}

_FALLBACK_KEYWORDS = {
    "rate_limit", "rate limit",
    "token", "context_length", "context length", "context window",
    "quota", "credit", "insufficient", "overloaded", "capacity",
    "exceeded", "too many", "too large",
}


def _doit_basculer(exc: Exception) -> bool:
    """Retourne True si cette erreur doit déclencher le basculement vers le fallback."""
    if isinstance(exc, openai.RateLimitError):
        return True
    if isinstance(exc, openai.APIStatusError):
        if exc.status_code in _FALLBACK_STATUS_CODES:
            return True
    msg = str(exc).lower()
    return any(kw in msg for kw in _FALLBACK_KEYWORDS)


# ── Wrapper fallback ───────────────────────────────────────────────────────

class _FallbackCompletions:
    """
    Wrapper autour de deux clients OpenAI-compatible.
    Essaie le provider primaire ; bascule vers le fallback sur erreur de quota/token.
    Chaque provider a son propre plafond max_tokens pour éviter les erreurs de crédits.
    """

    def __init__(
        self,
        primary_client: OpenAI,          primary_model: str,
        fallback_client: Optional[OpenAI], fallback_model: str,
        primary_max_tokens: int = 4096,
        fallback_max_tokens: int = 1500,
    ):
        self._primary              = primary_client
        self._primary_model        = primary_model
        self._primary_max_tokens   = primary_max_tokens
        self._fallback             = fallback_client
        self._fallback_model       = fallback_model
        self._fallback_max_tokens  = fallback_max_tokens

    def _cap_tokens(self, kwargs: dict, max_tokens: int) -> dict:
        """Plafonne max_tokens dans kwargs sans modifier l'original."""
        current = kwargs.get("max_tokens", max_tokens)
        if current > max_tokens:
            logger.info(f"max_tokens plafonné : {current} → {max_tokens}")
            return {**kwargs, "max_tokens": max_tokens}
        return kwargs

    def create(self, model: str = "", **kwargs) -> Any:
        # model passé par l'appelant est ignoré — on utilise le modèle configuré
        kwargs.pop("model", None)

        primary_kwargs = self._cap_tokens(kwargs, self._primary_max_tokens)
        try:
            logger.info(f"IA → primaire ({self._primary_model}, max_tokens={primary_kwargs.get('max_tokens')})")
            return self._primary.chat.completions.create(
                model=self._primary_model, **primary_kwargs
            )
        except Exception as exc:
            if self._fallback and _doit_basculer(exc):
                fallback_kwargs = self._cap_tokens(kwargs, self._fallback_max_tokens)
                logger.warning(
                    f"Provider primaire échoué [{type(exc).__name__}] : {exc} "
                    f"→ fallback ({self._fallback_model}, max_tokens={fallback_kwargs.get('max_tokens')})"
                )
                return self._fallback.chat.completions.create(
                    model=self._fallback_model, **fallback_kwargs
                )
            raise


class _ChatNamespace:
    def __init__(self, completions: _FallbackCompletions):
        self.completions = completions


class FallbackLLMClient:
    """
    Drop-in replacement pour openai.OpenAI avec logique de fallback intégrée.
    Expose client.chat.completions.create(...) — même interface que OpenAI SDK.
    """

    def __init__(
        self,
        primary_client: OpenAI,          primary_model: str,
        fallback_client: Optional[OpenAI] = None, fallback_model: str = "",
        primary_max_tokens: int = 4096,
        fallback_max_tokens: int = 1500,
    ):
        completions = _FallbackCompletions(
            primary_client, primary_model,
            fallback_client, fallback_model,
            primary_max_tokens, fallback_max_tokens,
        )
        self.chat = _ChatNamespace(completions)
        self._primary_model  = primary_model
        self._fallback_model = fallback_model
        self._has_fallback   = fallback_client is not None

    def __bool__(self) -> bool:
        return True


# ── Construction du client global ──────────────────────────────────────────

def _construire_client() -> Optional[FallbackLLMClient]:
    """Construit le client IA avec fallback selon les clés disponibles."""
    primary_client  = None
    fallback_client = None

    if NVIDIA_API_KEY:
        primary_client = OpenAI(
            api_key=NVIDIA_API_KEY,
            base_url="https://integrate.api.nvidia.com/v1",
        )

    if OPENROUTER_API_KEY:
        fallback_client = OpenAI(
            api_key=OPENROUTER_API_KEY,
            base_url="https://openrouter.ai/api/v1",
        )

    if primary_client:
        return FallbackLLMClient(
            primary_client,  NVIDIA_MODEL,
            fallback_client, OPENROUTER_MODEL,
            primary_max_tokens=NVIDIA_MAX_TOKENS,
            fallback_max_tokens=OPENROUTER_MAX_TOKENS,
        )
    if fallback_client:
        # Pas de clé NVIDIA → OpenRouter seul, sans fallback
        return FallbackLLMClient(
            fallback_client, OPENROUTER_MODEL,
            primary_max_tokens=OPENROUTER_MAX_TOKENS,
        )
    return None


# Client global partagé avec document_service.py
client = _construire_client()

# Modèle par défaut (pour affichage dans /health)
MODEL = NVIDIA_MODEL if NVIDIA_API_KEY else OPENROUTER_MODEL

# ── FastAPI ────────────────────────────────────────────────────────────────
app = FastAPI(
    title="SmartIntern AI — CV Extraction Service",
    description="Extraction, standardisation et scoring des CVs",
    version="2.1.0",
)
app.add_middleware(CORSMiddleware, allow_origins=["*"],
                   allow_methods=["*"], allow_headers=["*"])

ACCEPTED_TYPES = {
    "application/pdf": "pdf",
    "image/jpeg": "image", "image/jpg": "image",
    "image/png": "image", "image/webp": "image",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "docx",
    "application/octet-stream": "auto",
}
ACCEPTED_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png", ".webp", ".docx", ".doc"}


# ══════════════════════════════════════════════════════════════════════════
# MODELES PYDANTIC — alignes sur les entites Java
# alias_generator=to_camel : snake_case Python → camelCase JSON
# Compatible avec Jackson sans configuration supplementaire
# ══════════════════════════════════════════════════════════════════════════

class CamelModel(BaseModel):
    model_config = ConfigDict(alias_generator=to_camel, populate_by_name=True)


class Profil(CamelModel):
    nom: Optional[str] = None
    prenom: Optional[str] = None
    email: Optional[str] = None
    telephone: Optional[str] = None
    adresse: Optional[str] = None
    nationalite: Optional[str] = None
    date_naissance: Optional[str] = None
    titre_professionnel: Optional[str] = None
    resume: Optional[str] = None


class Experience(CamelModel):
    poste: str
    entreprise: str
    periode: Optional[str] = None
    date_debut: Optional[str] = None
    date_fin: Optional[str] = None
    description: Optional[str] = None
    lieu: Optional[str] = None


class Formation(CamelModel):
    diplome: str
    etablissement: str
    periode: Optional[str] = None
    date_debut: Optional[str] = None
    date_fin: Optional[str] = None
    specialite: Optional[str] = None
    lieu: Optional[str] = None


class Competences(CamelModel):
    techniques: list[str] = Field(default_factory=list)
    soft_skills: list[str] = Field(default_factory=list)   # → softSkills en JSON
    outils: list[str] = Field(default_factory=list)
    autres: list[str] = Field(default_factory=list)


class Langue(CamelModel):
    langue: str
    niveau: Optional[str] = None


def _dict_vers_str(item: object) -> str:
    """
    Convertit un dict renvoyé par un LLM en chaîne lisible.
    Ex: {"nom": "Photographie", "organisme": "Sophaera", "date": "2014"}
        → "Photographie — Sophaera (2014)"
    """
    if isinstance(item, str):
        return item
    if not isinstance(item, dict):
        return str(item)
    nom       = item.get("nom")       or item.get("name")        or item.get("titre")    or ""
    organisme = item.get("organisme") or item.get("organization") or item.get("organisation") \
                or item.get("etablissement") or item.get("entreprise") or ""
    date      = item.get("date")      or item.get("annee")        or item.get("year")    or ""
    desc      = item.get("description") or item.get("technologie") or item.get("technologies") or ""
    parts = [str(nom)]
    if organisme:
        parts.append(str(organisme))
    detail = []
    if date:
        detail.append(str(date))
    if desc and isinstance(desc, str) and len(desc) < 60:
        detail.append(desc)
    if detail:
        parts.append(f"({', '.join(detail)})")
    return " — ".join(filter(None, parts)) or str(item)


def _normaliser_liste_str(v: Any) -> list[str]:
    """BeforeValidator : convertit chaque élément dict → str avant validation Pydantic."""
    if not isinstance(v, list):
        return []
    return [_dict_vers_str(item) for item in v if item is not None]


StrList = Annotated[list[str], BeforeValidator(_normaliser_liste_str)]


class CvStandardise(CamelModel):
    """Miroir exact du DTO Java CvExtractionService.CvStandardise + scoreCompletude."""
    profil: Profil
    experiences: list[Experience] = Field(default_factory=list)
    formations: list[Formation] = Field(default_factory=list)
    competences: Competences = Field(default_factory=Competences)
    langues: list[Langue] = Field(default_factory=list)
    certifications: StrList = Field(default_factory=list)
    projets: StrList = Field(default_factory=list)
    interets: StrList = Field(default_factory=list)
    score_completude: float = Field(default=0.0)   # → scoreCompletude


# ── Document entities (Sprint 3 - generation de documents) ────────────────

class TypeDocument(CamelModel):
    nom: str
    description: Optional[str] = None


class Document(CamelModel):
    url_fichier: str
    numero_version: int = 1
    created_at: Optional[str] = None


# ── Score models ───────────────────────────────────────────────────────────

class ScoreSection(CamelModel):
    score: float
    max: float
    pourcentage: float
    manquant: list[str] = Field(default_factory=list)
    complet: bool


class ScoreResultat(CamelModel):
    score_global: float
    score_sur_100: float
    niveau: str
    details: dict
    recommandations: list[str]
    sections_completes: int
    sections_totales: int


class CvResponse(CamelModel):
    """
    Miroir exact du DTO Java CvExtractionService.CvResponse.
    Serialisation :
      cv_standardise      → cvStandardise
      texte_brut          → texteBrut
      nb_pages            → nbPages
      format_detecte      → formatDetecte
      methode_extraction  → methodeExtraction
    """
    success: bool
    cv_standardise: Optional[CvStandardise] = None
    texte_brut: Optional[str] = None
    nb_pages: int = 0
    format_detecte: str = ""
    methode_extraction: str = ""
    message: str = ""
    score: Optional[ScoreResultat] = None


# ══════════════════════════════════════════════════════════════════════════
# EXTRACTEURS
# ══════════════════════════════════════════════════════════════════════════

def extraire_depuis_pdf(chemin: str) -> tuple[str, int, str]:
    texte_pages, nb_pages = [], 0
    try:
        with pdfplumber.open(chemin) as pdf:
            nb_pages = len(pdf.pages)
            for page in pdf.pages:
                t = page.extract_text(layout=True) or page.extract_text() or ""
                texte_pages.append(t)
        texte = "\n\n--- PAGE ---\n\n".join(texte_pages).strip()
        if len(texte) >= 200:
            logger.info(f"PDF texte natif OK ({len(texte)} chars, {nb_pages} pages)")
            return texte, nb_pages, "pdf_text"
    except Exception as e:
        logger.warning(f"pdfplumber erreur: {e}")

    logger.info("PDF scanne detecte -> OCR")
    doc = fitz.open(chemin)
    nb_pages = len(doc)
    texte_pages_ocr = []
    for page in doc:
        mat = fitz.Matrix(200 / 72, 200 / 72)
        pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_img:
            pix.save(tmp_img.name)
            img = Image.open(tmp_img.name)
            texte_pages_ocr.append(pytesseract.image_to_string(img, lang="fra+eng"))
            os.unlink(tmp_img.name)
    texte = "\n\n--- PAGE ---\n\n".join(texte_pages_ocr).strip()
    return texte, nb_pages, "pdf_ocr"


def extraire_depuis_image(chemin: str) -> tuple[str, int, str]:
    img = Image.open(chemin)
    w, h = img.size
    if w < 1000:
        img = img.resize((int(1000), int(h * 1000 / w)), Image.LANCZOS)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    texte = pytesseract.image_to_string(img, lang="fra+eng")
    return texte.strip(), 1, "image_ocr"


def extraire_depuis_docx(chemin: str) -> tuple[str, int, str]:
    try:
        result = subprocess.run(
            ["pandoc", chemin, "-t", "plain", "--wrap=none"],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip(), 1, "docx_pandoc"
    except Exception as e:
        logger.warning(f"pandoc erreur: {e}")
    try:
        from docx import Document as DocxDoc
        doc = DocxDoc(chemin)
        texte = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return texte, 1, "docx_python"
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Impossible de lire le DOCX: {e}")


def detecter_format(filename: str, content_type: str) -> str:
    ext = Path(filename or "").suffix.lower()
    if ext in (".jpg", ".jpeg", ".png", ".webp"):
        return "image"
    if ext == ".pdf":
        return "pdf"
    if ext in (".docx", ".doc"):
        return "docx"
    return ACCEPTED_TYPES.get(content_type, "")


def extraire_texte(chemin: str, fmt: str) -> tuple[str, int, str]:
    if fmt == "pdf":    return extraire_depuis_pdf(chemin)
    if fmt == "image":  return extraire_depuis_image(chemin)
    if fmt == "docx":   return extraire_depuis_docx(chemin)
    raise HTTPException(status_code=400, detail=f"Format non supporte: {fmt}")


# ══════════════════════════════════════════════════════════════════════════
# ANALYSE IA
# ══════════════════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """Tu es un expert en analyse de CVs multilingues (français, anglais, arabe).
Extrait et structure les informations en JSON standardise.

REGLES :
1. Texte peut etre desordonne (OCR) — raisonne sur le sens
2. Distingue formations academiques VS certifications courtes
3. Competences TECHNIQUES != SOFT SKILLS != OUTILS
4. Ne JAMAIS inventer une info absente -> null
5. Dates : format lisible ("Janvier 2020 - Mars 2023")
6. Reponds UNIQUEMENT avec un JSON valide, SANS markdown, SANS texte avant ou apres
7. CRITIQUE — certifications, projets, interets : UNIQUEMENT des tableaux de CHAINES DE TEXTE simples.
   Format attendu  : ["Titre (Organisme, Annee)", "Autre titre"]
   FORMAT INTERDIT : [{"nom": "...", "date": "..."}, ...]  <- NE JAMAIS faire ca

Structure JSON exacte (respecter ces noms de cles) :
{
  "profil": {"nom": null, "prenom": null, "email": null, "telephone": null,
             "adresse": null, "nationalite": null, "date_naissance": null,
             "titre_professionnel": null, "resume": null},
  "experiences": [{"poste": "", "entreprise": "", "periode": null,
                   "date_debut": null, "date_fin": null, "description": null, "lieu": null}],
  "formations": [{"diplome": "", "etablissement": "", "periode": null,
                  "date_debut": null, "date_fin": null, "specialite": null, "lieu": null}],
  "competences": {"techniques": [], "soft_skills": [], "outils": [], "autres": []},
  "langues": [{"langue": "", "niveau": null}],
  "certifications": ["Exemple : Formation Python (Coursera, 2024)", "TOEIC 850 (2023)"],
  "projets": ["Exemple : Application mobile de gestion RH (React Native, 2024)"],
  "interets": ["Exemple : Photographie", "Musique"]
}"""


def nettoyer_json(texte: str) -> str:
    """Parsing JSON robuste par regex — remplace l'ancien split fragile."""
    texte = texte.strip()
    m = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", texte)
    if m:
        return m.group(1).strip()
    m = re.search(r"(\{[\s\S]*\})", texte)
    if m:
        return m.group(1).strip()
    return texte


def analyser_cv_avec_ia(texte_brut: str) -> CvStandardise:
    if not client:
        raise HTTPException(status_code=503,
            detail="Aucune clé API IA configurée. Renseignez NVIDIA_API_KEY et/ou OPENROUTER_API_KEY dans .env")
    try:
        response = client.chat.completions.create(
            model=MODEL,  # ignoré par FallbackLLMClient — utilise le modèle configuré
            max_tokens=16384,
            temperature=0.10,
            top_p=1.00,
            extra_body={"reasoning_effort": "high"},
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user",
                 "content": f"CV a analyser :\n\n{texte_brut}\n\nRetourne le JSON."}
            ]
        )
        json_propre = nettoyer_json(response.choices[0].message.content)
        data = json.loads(json_propre)

        # Normalisation préventive AVANT Pydantic : certains LLMs (Gemma, Mistral…)
        # retournent certifications/projets/interets sous forme de dicts structurés
        # {"nom": "...", "organisme": "...", "date": "..."} au lieu de strings simples.
        # On convertit ici sur le dict brut pour contourner tout comportement Pydantic v2.
        for _champ in ("certifications", "projets", "interets"):
            brut = data.get(_champ)
            if isinstance(brut, list):
                data[_champ] = [
                    _dict_vers_str(item) if not isinstance(item, str) else item
                    for item in brut
                    if item is not None
                ]

        return CvStandardise(**data)
    except json.JSONDecodeError as e:
        logger.error(f"JSON parse erreur: {e}")
        raise HTTPException(status_code=500, detail=f"Erreur structuration CV: {e}")
    except Exception as e:
        logger.error(f"Erreur API IA (tous providers épuisés): {e}")
        raise HTTPException(status_code=502, detail=f"Erreur API IA: {e}")


# ══════════════════════════════════════════════════════════════════════════
# CALCUL DU SCORE DE COMPLETUDE — 100 points / 6 sections
#
# Bareme :
#   Profil de base   : 25 pts
#   Formations       : 20 pts
#   Experiences      : 20 pts
#   Competences      : 20 pts
#   Langues          :  5 pts
#   Bonus            : 10 pts
# ══════════════════════════════════════════════════════════════════════════

def calculer_score(cv: CvStandardise) -> ScoreResultat:
    details = {}
    recommandations = []
    sections_completes = 0
    TOTAL = 6

    # ── 1. Profil (25 pts) ─────────────────────────────────────────────────
    p, s, mx, miss = cv.profil, 0.0, 25.0, []

    if p.nom and p.prenom:        s += 5
    else:
        miss.append("Nom et prenom")
        recommandations.append("Indiquez votre nom complet en haut du CV")

    if p.email:                   s += 4
    else:
        miss.append("Email")
        recommandations.append("Ajoutez votre adresse email professionnelle")

    if p.telephone:               s += 3
    else:
        miss.append("Telephone")
        recommandations.append("Ajoutez votre numero de telephone")

    if p.titre_professionnel:     s += 5
    else:
        miss.append("Titre professionnel")
        recommandations.append("Ajoutez un titre professionnel (ex: Etudiant en Genie Logiciel)")

    if p.resume and len(p.resume) >= 50:
        s += 5
    elif p.resume:
        s += 2; miss.append("Resume trop court")
        recommandations.append("Developpez votre resume (minimum 50 caracteres)")
    else:
        miss.append("Resume / objectif professionnel")
        recommandations.append("Ajoutez un resume decrivant votre profil et objectifs")

    if p.adresse:                 s += 3
    else:
        miss.append("Localisation")
        recommandations.append("Indiquez votre ville ou region")

    ok = len(miss) == 0
    if ok: sections_completes += 1
    details["profil"] = ScoreSection(score=s, max=mx,
        pourcentage=round(s/mx*100,1), manquant=miss, complet=ok
    ).model_dump(by_alias=True)

    # ── 2. Formations (20 pts) ─────────────────────────────────────────────
    s, mx, miss = 0.0, 20.0, []

    if cv.formations:
        s += 10
        f = cv.formations[0]
        if f.specialite:          s += 3
        else:
            miss.append("Specialite"); recommandations.append("Precisez votre specialite pour chaque formation")
        if f.date_debut or f.periode: s += 3
        else:
            miss.append("Dates de formation"); recommandations.append("Ajoutez les dates de vos formations")
        if f.lieu:                s += 2
        else:
            miss.append("Lieu")
        if len(cv.formations) >= 2: s += 2
    else:
        miss.append("Aucune formation")
        recommandations.append("Ajoutez vos formations (diplome, etablissement, dates, specialite)")

    ok = s >= mx * 0.8
    if ok: sections_completes += 1
    details["formations"] = ScoreSection(score=s, max=mx,
        pourcentage=round(s/mx*100,1), manquant=miss, complet=ok
    ).model_dump(by_alias=True)

    # ── 3. Experiences (20 pts) ────────────────────────────────────────────
    s, mx, miss = 0.0, 20.0, []

    if cv.experiences:
        s += 8
        exp = cv.experiences[0]
        desc_ok = [e for e in cv.experiences if e.description and len(e.description) >= 30]
        if desc_ok:               s += 6
        else:
            miss.append("Descriptions insuffisantes")
            recommandations.append("Decrivez vos missions et realisations (minimum 30 caracteres)")
        if exp.date_debut or exp.periode: s += 3
        else:
            miss.append("Dates des experiences"); recommandations.append("Ajoutez les dates de chaque experience")
        if exp.lieu:              s += 1
        if len(cv.experiences) >= 2: s += 2
    else:
        miss.append("Aucune experience")
        recommandations.append("Ajoutez vos stages, projets academiques ou experiences professionnelles")

    ok = s >= mx * 0.7
    if ok: sections_completes += 1
    details["experiences"] = ScoreSection(score=s, max=mx,
        pourcentage=round(s/mx*100,1), manquant=miss, complet=ok
    ).model_dump(by_alias=True)

    # ── 4. Competences (20 pts) ────────────────────────────────────────────
    s, mx, miss = 0.0, 20.0, []
    c = cv.competences

    nb_t = len(c.techniques)
    if nb_t >= 5:                 s += 8
    elif nb_t >= 3:               s += 5
    elif nb_t >= 1:
        s += 2; miss.append(f"Peu de competences techniques ({nb_t})")
        recommandations.append("Listez au moins 5 competences techniques (langages, frameworks, BDD)")
    else:
        miss.append("Aucune competence technique")
        recommandations.append("Ajoutez vos competences techniques : langages, frameworks, technologies")

    nb_o = len(c.outils)
    if nb_o >= 3:                 s += 5
    elif nb_o >= 1:
        s += 3; miss.append("Peu d'outils renseignes")
        recommandations.append("Ajoutez les outils que vous maitrisez (IDE, Git, Docker, etc.)")
    else:
        miss.append("Aucun outil renseigne")
        recommandations.append("Listez vos outils et logiciels")

    nb_s = len(c.soft_skills)
    if nb_s >= 3:                 s += 5
    elif nb_s >= 1:
        s += 3; miss.append("Peu de soft skills")
        recommandations.append("Ajoutez vos competences comportementales (equipe, communication)")
    else:
        miss.append("Aucun soft skill")
        recommandations.append("Ajoutez des competences transversales : autonomie, rigueur, esprit equipe")

    if nb_t + nb_o + nb_s >= 10: s += 2

    ok = s >= mx * 0.7
    if ok: sections_completes += 1
    details["competences"] = ScoreSection(score=s, max=mx,
        pourcentage=round(s/mx*100,1), manquant=miss, complet=ok
    ).model_dump(by_alias=True)

    # ── 5. Langues (5 pts) ────────────────────────────────────────────────
    s, mx, miss = 0.0, 5.0, []

    if cv.langues:
        s += 3
        if any(l.niveau for l in cv.langues): s += 2
        else:
            miss.append("Niveau de langue non precise")
            recommandations.append("Precisez votre niveau pour chaque langue (Debutant / Intermediaire / Avance / Natif)")
    else:
        miss.append("Aucune langue")
        recommandations.append("Ajoutez les langues que vous parlez avec votre niveau")

    ok = s >= mx * 0.8
    if ok: sections_completes += 1
    details["langues"] = ScoreSection(score=s, max=mx,
        pourcentage=round(s/mx*100,1), manquant=miss, complet=ok
    ).model_dump(by_alias=True)

    # ── 6. Bonus (10 pts) ────────────────────────────────────────────────
    s, mx, miss = 0.0, 10.0, []

    if cv.certifications:         s += 3
    else:
        miss.append("Aucune certification")
        recommandations.append("Ajoutez vos certifications (MOOC, diplomes complementaires)")

    if cv.projets:                s += 4
    else:
        miss.append("Aucun projet")
        recommandations.append("Decrivez vos projets personnels ou academiques (GitHub, hackathons)")

    if cv.interets:               s += 3
    else:
        miss.append("Aucun interet")
        recommandations.append("Ajoutez quelques centres d'interet pour humaniser votre candidature")

    ok = s >= mx * 0.7
    if ok: sections_completes += 1
    details["bonus"] = ScoreSection(score=s, max=mx,
        pourcentage=round(s/mx*100,1), manquant=miss, complet=ok
    ).model_dump(by_alias=True)

    # ── Score global ───────────────────────────────────────────────────────
    score_total = sum(details[k]["score"] for k in details)
    max_total   = sum(details[k]["max"]   for k in details)
    sur_100 = round(score_total / max_total * 100, 1)
    niveau  = ("EXCELLENT" if sur_100 >= 85 else
               "BON"       if sur_100 >= 65 else
               "MOYEN"     if sur_100 >= 40 else "INSUFFISANT")

    return ScoreResultat(
        score_global=round(score_total, 1),
        score_sur_100=sur_100,
        niveau=niveau,
        details=details,
        recommandations=recommandations,
        sections_completes=sections_completes,
        sections_totales=TOTAL,
    )


# ══════════════════════════════════════════════════════════════════════════
# ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════

@app.get("/health")
async def health_check():
    providers = []
    if NVIDIA_API_KEY:
        providers.append({
            "nom": "nvidia_nim",
            "role": "primaire",
            "modele": NVIDIA_MODEL,
            "configure": True,
        })
    if OPENROUTER_API_KEY:
        providers.append({
            "nom": "openrouter",
            "role": "fallback" if NVIDIA_API_KEY else "primaire",
            "modele": OPENROUTER_MODEL,
            "configure": True,
        })
    return {
        "status": "ok",
        "service": "cv-extraction",
        "version": "2.1.0",
        "ia_disponible": client is not None,
        "fallback_actif": bool(client and client._has_fallback) if client else False,
        "providers": providers,
        "formats_supportes": ["pdf", "pdf_ocr", "jpg", "png", "webp", "docx"],
        "timestamp": datetime.now().isoformat(),
    }


@app.post("/extract")
async def extraire_cv(file: UploadFile = File(...)):
    """
    Extraction + structuration IA + score de completude.
    Retourne CvResponse en camelCase (compatible Java/Jackson).
    """
    filename = file.filename or ""
    content_type = file.content_type or ""

    format_fichier = detecter_format(filename, content_type)
    if not format_fichier or format_fichier == "auto":
        ext = Path(filename).suffix.lower()
        if ext not in ACCEPTED_EXTENSIONS:
            raise HTTPException(status_code=400,
                detail="Format non supporte. Acceptes : PDF, JPG, PNG, WEBP, DOCX, DOC")
        format_fichier = detecter_format(filename, "")

    contenu = await file.read()
    if len(contenu) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Fichier trop volumineux (max 10MB)")

    logger.info(f"Traitement: {filename} | format={format_fichier} | {len(contenu)} bytes")

    ext = Path(filename).suffix or (".pdf" if format_fichier == "pdf" else ".tmp")
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(contenu)
        chemin_tmp = tmp.name

    try:
        texte_brut, nb_pages, methode = extraire_texte(chemin_tmp, format_fichier)

        if len(texte_brut.strip()) < 50:
            return CvResponse(
                success=False, nb_pages=nb_pages,
                format_detecte=format_fichier, methode_extraction=methode,
                message="Impossible d'extraire le texte (fichier corrompu ou protege)"
            ).model_dump(by_alias=True)

        cv = analyser_cv_avec_ia(texte_brut)
        score = calculer_score(cv)
        cv.score_completude = score.score_sur_100

        logger.info(f"CV traite: {nb_pages}p | {methode} | score={score.score_sur_100}% ({score.niveau})")

        return CvResponse(
            success=True,
            cv_standardise=cv,
            texte_brut=texte_brut,
            nb_pages=nb_pages,
            format_detecte=format_fichier,
            methode_extraction=methode,
            message="CV extrait, structure et score avec succes",
            score=score,
        ).model_dump(by_alias=True)

    finally:
        if os.path.exists(chemin_tmp):
            os.unlink(chemin_tmp)


@app.post("/extract/text-only")
async def extraire_texte_seulement(file: UploadFile = File(...)):
    """Extraction texte brut uniquement, sans analyse IA ni scoring."""
    filename = file.filename or ""
    format_fichier = detecter_format(filename, file.content_type or "")
    contenu = await file.read()
    ext = Path(filename).suffix or ".tmp"
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(contenu)
        chemin_tmp = tmp.name
    try:
        texte_brut, nb_pages, methode = extraire_texte(chemin_tmp, format_fichier)
        return CvResponse(
            success=True, texte_brut=texte_brut, nb_pages=nb_pages,
            format_detecte=format_fichier, methode_extraction=methode,
            message="Texte extrait avec succes"
        ).model_dump(by_alias=True)
    finally:
        if os.path.exists(chemin_tmp):
            os.unlink(chemin_tmp)


@app.post("/score")
async def scorer_cv(cv: CvStandardise):
    """
    Calcule le score de completude d'un CV deja structure.
    Utilise quand l'etudiant modifie son profil manuellement.
    """
    score = calculer_score(cv)
    return score.model_dump(by_alias=True)


# ══════════════════════════════════════════════════════════════════════════
# INTEGRATION MODULE DOCUMENTS
# ══════════════════════════════════════════════════════════════════════════

from fastapi import Form
from fastapi.responses import FileResponse
from document_service import DocumentService, est_expire, verifier_signature

# Instance du service document (partage le client Anthropic)
doc_service = DocumentService(anthropic_client=client)


# ── Modeles Pydantic — Documents ─────────────────────────────────────────

class TypeDocumentModel(CamelModel):
    """Miroir de l'entite Java TypeDocument."""
    nom: str
    description: Optional[str] = None


class ModeleDocumentInfo(CamelModel):
    """Informations sur un modele de document cree."""
    modele_id: str
    nom_modele: str
    type_document: str
    titre_document: str
    langue: str
    chemin_modele: str
    chemin_header: Optional[str] = None
    chemin_footer: Optional[str] = None
    chemin_script: str
    champs_dynamiques: list = Field(default_factory=list)
    sections: list = Field(default_factory=list)
    mise_en_page: dict = Field(default_factory=dict)
    duree_validite_jours: int = 365
    version: int = 1
    date_creation: str = ""


class DemandeGenerationDocument(CamelModel):
    """
    Requete de generation d'un document.
    Envoye par le backend Java avec toutes les donnees profil.
    """
    modele_id: str
    chemin_script: str
    chemin_header: Optional[str] = None
    chemin_footer: Optional[str] = None
    type_document: str
    donnees_profil: dict = Field(default_factory=dict)
    nom_etablissement: str = "SmartIntern"
    duree_validite_jours: int = 365


class DemandeRegenerationDocument(DemandeGenerationDocument):
    doc_uuid_original: str


class DocumentGenereInfo(CamelModel):
    """
    Miroir de l'entite Java DocumentGenere.
    Retourne les metadonnees au backend Java pour persistance MySQL.
    """
    doc_uuid: str
    modele_id: str
    type_document: str
    chemin_fichier: str
    url_fichier: str
    url_verification: str
    date_generation: str
    date_expiration: str
    signature: str
    statut: str = "VALIDE"
    taille_octets: int = 0
    nom_etablissement: str = ""
    donnees_profil_utilises: list = Field(default_factory=list)
    doc_uuid_original: Optional[str] = None
    regeneration: bool = False


class ResultatVerification(CamelModel):
    """Resultat de la verification d'un document via QR code."""
    doc_uuid: str
    statut: str
    valide: bool
    message: str
    type_document: str = ""
    nom_etablissement: str = ""
    date_generation: str = ""
    date_expiration: str = ""
    signature_valide: bool = False
    expire: bool = False
    fichier_existe: bool = False
    verifie_le: str = ""


# ── ENDPOINTS — Modeles de documents ─────────────────────────────────────

@app.post("/modeles/creer")
async def creer_modele_document(
    modele_id: str = Form(...),
    nom_modele: str = Form(...),
    type_document: str = Form(...),
    duree_validite_jours: int = Form(default=365),
    fichier_modele: UploadFile = File(...),
    fichier_header: Optional[UploadFile] = File(default=None),
    fichier_footer: Optional[UploadFile] = File(default=None),
):
    """
    Cree un nouveau modele de document.

    Processus :
      1. Sauvegarde header PNG/JPG (optionnel)
      2. Sauvegarde footer PNG/JPG (optionnel)
      3. Sauvegarde modele PDF ou DOCX
      4. Analyse IA -> extraction champs dynamiques
      5. Generation automatique script Python
      6. Retourne metadonnees completes pour persistance Java/MySQL

    Formats acceptes :
      - fichier_modele : PDF, DOCX, DOC
      - fichier_header / fichier_footer : PNG, JPG, JPEG
    """
    if not client:
        raise HTTPException(status_code=503,
            detail="Aucune clé API IA configurée. Renseignez NVIDIA_API_KEY et/ou OPENROUTER_API_KEY dans .env")

    nom_modele_fichier = fichier_modele.filename or "modele.pdf"
    ext = Path(nom_modele_fichier).suffix.lower()
    if ext not in (".pdf", ".docx", ".doc"):
        raise HTTPException(status_code=400,
            detail="Modele : seuls PDF, DOCX, DOC sont acceptes")

    contenu_modele = await fichier_modele.read()
    if len(contenu_modele) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Modele trop volumineux (max 20MB)")

    contenu_header, nom_header = None, None
    if fichier_header and fichier_header.filename:
        ext_h = Path(fichier_header.filename).suffix.lower()
        if ext_h not in (".png", ".jpg", ".jpeg"):
            raise HTTPException(status_code=400,
                detail="Header : seuls PNG, JPG sont acceptes")
        contenu_header = await fichier_header.read()
        nom_header = fichier_header.filename

    contenu_footer, nom_footer = None, None
    if fichier_footer and fichier_footer.filename:
        ext_f = Path(fichier_footer.filename).suffix.lower()
        if ext_f not in (".png", ".jpg", ".jpeg"):
            raise HTTPException(status_code=400,
                detail="Footer : seuls PNG, JPG sont acceptes")
        contenu_footer = await fichier_footer.read()
        nom_footer = fichier_footer.filename

    logger.info(f"Creation modele: {nom_modele} | type={type_document} | id={modele_id}")

    try:
        resultat = doc_service.creer_modele(
            modele_id=modele_id,
            nom_modele=nom_modele,
            type_document=type_document,
            contenu_modele=contenu_modele,
            nom_fichier_modele=nom_modele_fichier,
            contenu_header=contenu_header,
            nom_fichier_header=nom_header,
            contenu_footer=contenu_footer,
            nom_fichier_footer=nom_footer,
            duree_validite_jours=duree_validite_jours,
        )
        return {
            "success": True,
            "message": f"Modele '{nom_modele}' cree avec succes",
            "modele": resultat,
            "nb_champs_detectes": len(resultat.get("champs_dynamiques", [])),
            "type_detecte": resultat.get("type_document"),
        }
    except Exception as e:
        logger.error(f"Erreur creation modele: {e}")
        raise HTTPException(status_code=500, detail=f"Erreur creation modele: {str(e)}")


@app.get("/modeles/{modele_id}/champs")
async def obtenir_champs_modele(modele_id: str):
    """
    Retourne les champs dynamiques d'un modele (pour affichage dans le frontend).
    Le backend Java stocke ces infos mais peut interroger le microservice.
    """
    chemin_script = Path("uploads/scripts") / f"{modele_id}_script.py"
    if not chemin_script.exists():
        raise HTTPException(status_code=404,
            detail=f"Modele {modele_id} introuvable")
    return {"modele_id": modele_id, "script_existe": True,
            "chemin_script": str(chemin_script)}


# ── ENDPOINTS — Generation de documents ──────────────────────────────────

@app.post("/documents/generer")
async def generer_document(demande: DemandeGenerationDocument):
    """
    Genere un document PDF personnalise pour un utilisateur.

    Appele par le backend Java avec les donnees profil de l'utilisateur.
    Le backend Java est responsable de persister le DocumentGenere en MySQL.

    Corps JSON attendu :
    {
      "modeleId": "uuid-du-modele",
      "cheminScript": "uploads/scripts/uuid_script.py",
      "cheminHeader": "uploads/headers/uuid_header.png",
      "cheminFooter": "uploads/footers/uuid_footer.png",
      "typeDocument": "convention_stage",
      "donneesProfil": {
        "nom": "Dupont", "prenom": "Jean",
        "email": "jean.dupont@etudiant.fr",
        "filiere": "Informatique", "classe": "3ING",
        "nom_entreprise": "TechCorp", "adresse_entreprise": "Paris",
        "date_debut_stage": "2024-06-01",
        "date_fin_stage": "2024-08-31"
      },
      "nomEtablissement": "ITEAM University",
      "dureeValiditeJours": 365
    }
    """
    logger.info(
        f"Generation document: modele={demande.modele_id} | "
        f"type={demande.type_document}"
    )
    try:
        resultat = doc_service.generer_document(
            modele_id=demande.modele_id,
            chemin_script=demande.chemin_script,
            chemin_header=demande.chemin_header,
            chemin_footer=demande.chemin_footer,
            type_document=demande.type_document,
            donnees_profil=demande.donnees_profil,
            nom_etablissement=demande.nom_etablissement,
            duree_validite_jours=demande.duree_validite_jours,
        )
        return {
            "success": True,
            "message": "Document genere avec succes",
            "document": resultat,
        }
    except Exception as e:
        logger.error(f"Erreur generation document: {e}")
        raise HTTPException(status_code=500,
            detail=f"Erreur generation document: {str(e)}")


@app.post("/documents/regenerer")
async def regenerer_document(demande: DemandeRegenerationDocument):
    """
    Regenere un document existant (nouvelle version).
    Conserve la reference au document original.
    """
    logger.info(
        f"Regeneration: original={demande.doc_uuid_original} | "
        f"modele={demande.modele_id}"
    )
    try:
        resultat = doc_service.regenerer_document(
            doc_uuid_original=demande.doc_uuid_original,
            modele_id=demande.modele_id,
            chemin_script=demande.chemin_script,
            chemin_header=demande.chemin_header,
            chemin_footer=demande.chemin_footer,
            type_document=demande.type_document,
            donnees_profil=demande.donnees_profil,
            nom_etablissement=demande.nom_etablissement,
            duree_validite_jours=demande.duree_validite_jours,
        )
        return {"success": True, "message": "Document regenere", "document": resultat}
    except Exception as e:
        raise HTTPException(status_code=500,
            detail=f"Erreur regeneration: {str(e)}")


@app.get("/documents/telecharger/{doc_uuid}")
async def telecharger_document(doc_uuid: str):
    """
    Telecharge le PDF d'un document genere.
    Endpoint appele par le frontend ou le backend Java.
    """
    chemin = Path("uploads/documents") / f"{doc_uuid}.pdf"
    if not chemin.exists():
        raise HTTPException(status_code=404,
            detail=f"Document {doc_uuid} introuvable")
    return FileResponse(
        str(chemin),
        media_type="application/pdf",
        filename=f"document_{doc_uuid[:8]}.pdf",
        headers={"Content-Disposition": f'attachment; filename="document_{doc_uuid[:8]}.pdf"'},
    )


# ── ENDPOINTS — Verification QR code ─────────────────────────────────────

@app.get("/documents/verifier/{doc_uuid}")
async def verifier_document_qr(
    doc_uuid: str,
    date_generation: str = "",
    date_expiration: str = "",
    signature: str = "",
    type_document: str = "",
    nom_etablissement: str = "",
):
    """
    Endpoint de verification d'authenticite scanne depuis le QR code.
    Accessible publiquement (pas d'authentification requise).

    Retourne :
      - statut : VALIDE | EXPIRE | FALSIFIE | INTROUVABLE
      - Informations du document si valide
      - Message clair si expire ou invalide
    """
    resultat = doc_service.verifier_document(
        doc_uuid=doc_uuid,
        date_generation=date_generation,
        date_expiration=date_expiration,
        signature=signature,
        type_document=type_document,
        nom_etablissement=nom_etablissement,
    )
    # Code HTTP selon le statut
    status_code = 200
    if resultat["statut"] == "EXPIRE":
        status_code = 410  # Gone
    elif resultat["statut"] in ("FALSIFIE", "INTROUVABLE"):
        status_code = 404
    from fastapi.responses import JSONResponse
    return JSONResponse(content=resultat, status_code=status_code)


@app.post("/documents/verifier")
async def verifier_document_post(donnees: dict):
    """
    Verification par POST (depuis le backend Java apres scan QR).
    Meme logique que GET mais accepte un JSON complet.
    """
    resultat = doc_service.verifier_document(
        doc_uuid=donnees.get("docUuid", ""),
        date_generation=donnees.get("dateGeneration", ""),
        date_expiration=donnees.get("dateExpiration", ""),
        signature=donnees.get("signature", ""),
        type_document=donnees.get("typeDocument", ""),
        nom_etablissement=donnees.get("nomEtablissement", ""),
    )
    return resultat


# ── ENDPOINT — Types de documents (reference) ─────────────────────────────

@app.get("/types-documents")
async def lister_types_documents():
    """
    Retourne la liste des types de documents supportes.
    Utilise comme reference pour la creation de modeles.
    """
    return {
        "types": [
            {"code": "convention_stage",        "label": "Convention de stage",           "categorie": "stage"},
            {"code": "attestation_stage",        "label": "Attestation de stage",          "categorie": "stage"},
            {"code": "lettre_recommandation",    "label": "Lettre de recommandation",      "categorie": "academique"},
            {"code": "attestation_inscription",  "label": "Attestation d'inscription",     "categorie": "academique"},
            {"code": "releve_notes",             "label": "Relevé de notes",               "categorie": "academique"},
            {"code": "demande_stage",            "label": "Demande de stage",              "categorie": "stage"},
            {"code": "rapport_stage",            "label": "Rapport de stage",              "categorie": "stage"},
            {"code": "fiche_evaluation",         "label": "Fiche d'évaluation",            "categorie": "evaluation"},
            {"code": "certificat_formation",     "label": "Certificat de formation",       "categorie": "formation"},
            {"code": "autre",                    "label": "Autre document",                "categorie": "autre"},
        ]
    }
